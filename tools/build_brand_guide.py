from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\www\mattos")
KIT = ROOT / "brand-kit"
IMG = KIT / "imagens"
OUT = KIT / "Guia_de_Marca_Mattos_Solucoes_Digitais.docx"

NAVY = "0F172A"
DEEP = "071A3A"
BLUE = "2563EB"
TEAL = "14B8A6"
SLATE = "64748B"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "111827"


def set_font(run, name="Calibri", size=11, color=BLACK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("MATTOS • GUIA DE MARCA   |   ")
    set_font(run, size=8.5, color=SLATE, bold=True)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, NAVY),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, DEEP),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "MATTOS SOLUÇÕES DIGITAIS"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.runs[0], size=8.5, color=SLATE, bold=True)
    add_page_number(section.footer.paragraphs[0])


def kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    set_font(r, size=9, color=TEAL, bold=True)
    r.font.letter_spacing = Pt(1.5)
    return p


def title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_font(r, size=25, color=NAVY, bold=True)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(14)
        r2 = p2.add_run(subtitle)
        set_font(r2, size=12.5, color=SLATE)


def add_callout(doc, label, text, fill=PALE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label.upper() + "  ")
    set_font(r, size=9, color=TEAL, bold=True)
    r2 = p.add_run(text)
    set_font(r2, size=10.5, color=NAVY, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_image(doc, path, width, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption or path.stem)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        rr = cp.add_run(caption)
        set_font(rr, size=8.5, color=SLATE, italic=True)


def add_two_images(doc, left, right, captions, widths=(3.0, 3.0)):
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [4680, 4680], indent_dxa=0)
    for idx, (img, width, caption) in enumerate(zip((left, right), widths, captions)):
        cell = table.cell(0, idx)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = p.add_run().add_picture(str(img), width=Inches(width))
        shape._inline.docPr.set("descr", caption)
        cp = table.cell(1, idx).paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = cp.add_run(caption)
        set_font(rr, size=8.5, color=SLATE, italic=True)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=60, bottom=60, start=80, end=80)


doc = Document()
style_doc(doc)

# Capa: editorial_cover com override de cores da marca.
for _ in range(3):
    doc.add_paragraph()
kicker(doc, "Manual de identidade visual")
add_image(doc, IMG / "01-logo-principal-fundo-claro.png", 6.25)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after = Pt(6)
r = p.add_run("GUIA DE MARCA E APLICAÇÕES")
set_font(r, size=22, color=NAVY, bold=True)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Sistema visual atualizado • Agosto de 2026")
set_font(r2, size=10.5, color=SLATE)
add_callout(doc, "Atualização principal", "A assinatura verbal passa a ser “Soluções digitais”, substituindo “Digital Solutions”.")

doc.add_page_break()
kicker(doc, "01 • Fundamentos")
title(doc, "Essência da marca", "Uma identidade de tecnologia com foco em clareza, crescimento e presença comercial.")
add_callout(doc, "Conceito central", "Presença digital que transforma visibilidade em oportunidades de negócio.")
doc.add_heading("Ideia visual", level=2)
doc.add_paragraph("O monograma combina a inicial M com formas ascendentes. A diagonal azul sugere movimento, evolução e tecnologia; o pilar verde-água reforça construção, continuidade e resultado.")
doc.add_heading("Personalidade", level=2)
add_bullets(doc, [
    "Profissional e segura, sem parecer distante.",
    "Tecnológica, porém acessível para empresas de diferentes portes.",
    "Direta, orientada a resultado e visualmente organizada.",
    "Contemporânea, com contraste forte e acentos de energia.",
])
doc.add_heading("Assinatura verbal", level=2)
doc.add_paragraph("Nome: Mattos. Descritor: Soluções digitais. Conceito de apoio: Presença digital que gera negócios.")
add_callout(doc, "Diretriz de escrita", "Em textos corridos, use “Soluções digitais”. Em peças gráficas, a caixa alta é permitida: “SOLUÇÕES DIGITAIS”.", fill=LIGHT)

doc.add_page_break()
kicker(doc, "02 • Sistema de logos")
title(doc, "Arquitetura da marca", "Escolha a assinatura conforme espaço, fundo e nível de reconhecimento necessário.")
add_image(doc, IMG / "01-logo-principal-fundo-claro.png", 6.35, "Logo principal — versão institucional completa")
add_two_images(doc, IMG / "03-logo-horizontal-fundo-claro.png", IMG / "05-monograma.png", ("Logo horizontal — cabeçalhos e assinaturas", "Monograma — avatar, favicon e espaços compactos"), widths=(3.35, 1.6))
doc.add_heading("Área de proteção", level=2)
doc.add_paragraph("Mantenha ao redor da marca uma margem mínima equivalente à largura do pilar verde-água do monograma. Não encoste textos, bordas, fotos ou outros símbolos nessa área.")
doc.add_heading("Tamanho mínimo recomendado", level=2)
add_bullets(doc, ["Logo principal: 240 px em meios digitais.", "Logo horizontal: 180 px em meios digitais.", "Monograma: 32 px; abaixo disso, simplifique efeitos e sombras."])

doc.add_page_break()
kicker(doc, "03 • Cores e tipografia")
title(doc, "Sistema visual", "A paleta equilibra confiança, dinamismo e inovação.")
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
mark_header_row(table.rows[0])
headers = ["COR", "HEX", "PAPEL", "USO PRINCIPAL"]
for i, h in enumerate(headers):
    shade(table.cell(0, i), NAVY)
    rr = table.cell(0, i).paragraphs[0].add_run(h)
    set_font(rr, size=8.5, color=WHITE, bold=True)
colors = [
    ("Deep Navy", "#071A3A", "Base", "Fundos e contraste"),
    ("Electric Blue", "#2563EB", "Ação", "Destaques e tecnologia"),
    ("Teal", "#14B8A6", "Crescimento", "Resultados e chamadas"),
    ("Slate Gray", "#64748B", "Apoio", "Textos secundários"),
    ("Off White", "#F8FAFC", "Respiro", "Fundos claros"),
]
for idx, row in enumerate(colors):
    cells = table.add_row().cells
    for i, value in enumerate(row):
        rr = cells[i].paragraphs[0].add_run(value)
        set_font(rr, size=9.5, color=BLACK, bold=(i == 0))
        if idx % 2:
            shade(cells[i], PALE)
set_table_geometry(table, [1800, 1500, 1800, 4260])
doc.add_heading("Tipografia", level=2)
doc.add_paragraph("Poppins Bold é a referência para títulos e destaques. Inter Regular é indicada para textos, informações e interfaces. Quando essas fontes não estiverem disponíveis, use Arial Bold e Arial como substitutas seguras.")
add_callout(doc, "Exemplo de hierarquia", "Sua empresa. Mais visível. Mais clientes. Mais resultados.")
doc.add_heading("Proporção de uso", level=2)
add_bullets(doc, ["60% fundos claros ou off-white.", "25% azul-marinho para peças de alto contraste.", "10% azul elétrico para ação e ênfase.", "5% teal para resultados, crescimento e pontos de atenção."])

doc.add_page_break()
kicker(doc, "04 • Redes sociais")
title(doc, "Aplicações digitais", "Arquivos dimensionados para perfil e feed do Instagram.")
add_two_images(doc, IMG / "06-avatar-instagram.png", IMG / "07-assinatura-post-instagram.png", ("Avatar institucional — 1080 × 1080 px", "Assinatura para feed — 1080 × 1350 px"), widths=(2.75, 2.55))
doc.add_heading("Boas práticas", level=2)
add_bullets(doc, [
    "Use o monograma no avatar; ele permanece reconhecível em miniatura.",
    "Em posts, reserve uma faixa de respiro ao redor do logo e evite fundos visualmente ruidosos.",
    "Prefira o logo claro sobre fundos escuros e o logo azul-marinho sobre fundos brancos.",
    "Não altere cores, proporções, espaçamentos ou a posição relativa do descritor.",
])
add_callout(doc, "Arquivo recomendado", "Para sobrepor a marca em posts, use a versão PNG transparente. Para máxima qualidade, preserve também o SVG.", fill=LIGHT)

doc.add_page_break()
kicker(doc, "05 • Cartão de visita")
title(doc, "Aplicação impressa", "Modelo horizontal em 3,5 × 2 polegadas, exportado a 300 dpi.")
add_image(doc, IMG / "08-cartao-visita-frente.png", 5.25, "Frente — assinatura institucional em fundo azul-marinho")
add_image(doc, IMG / "09-cartao-visita-verso-editavel.png", 5.25, "Verso — dados profissionais e canais de contato")
add_callout(doc, "Antes de imprimir", "Substitua “Seu Nome”, cargo, telefone e demais canais. Confirme domínio e usuário do Instagram. Solicite à gráfica sangria de 3 mm e mantenha textos a pelo menos 4 mm do corte.")

doc.add_page_break()
kicker(doc, "06 • Governança e arquivos")
title(doc, "Como reutilizar o kit", "Mantenha os arquivos-fonte e derive novas peças sempre a partir deles.")
doc.add_heading("O que não fazer", level=2)
add_bullets(doc, [
    "Não esticar, inclinar ou redesenhar o monograma.",
    "Não substituir o azul e o teal por cores aleatórias.",
    "Não aplicar sombras pesadas, contornos ou efeitos 3D sobre o logo.",
    "Não voltar a usar “Digital Solutions” em novas peças.",
    "Não posicionar a versão escura sobre fundos de baixo contraste.",
])
doc.add_heading("Índice de entregáveis", level=2)
table = doc.add_table(rows=1, cols=3)
mark_header_row(table.rows[0])
for i, h in enumerate(("PASTA", "FORMATO", "FINALIDADE")):
    shade(table.cell(0, i), NAVY)
    rr = table.cell(0, i).paragraphs[0].add_run(h)
    set_font(rr, size=8.5, color=WHITE, bold=True)
inventory = [
    ("imagens/01–05", "PNG", "Logos para fundo claro, escuro e transparente"),
    ("imagens/06–07", "PNG", "Avatar e assinatura para Instagram"),
    ("imagens/08–09", "PNG", "Cartão de visita — frente e verso"),
    ("vetores/", "SVG", "Fontes editáveis e escaláveis"),
    ("Guia de Marca", "DOCX", "Conceito, padrões e instruções de uso"),
]
for row in inventory:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        rr = cells[i].paragraphs[0].add_run(value)
        set_font(rr, size=9.5, color=BLACK, bold=(i == 0))
set_table_geometry(table, [2100, 1500, 5760])
doc.add_heading("Fluxo recomendado", level=2)
add_bullets(doc, [
    "Escolha o formato da peça e defina o fundo.",
    "Selecione a versão correta do logo.",
    "Aplique a paleta e a tipografia conforme este guia.",
    "Revise contraste, margens e legibilidade em tamanho real.",
    "Exporte em PNG para redes sociais e mantenha o SVG como matriz.",
])
add_callout(doc, "Documento vivo", "Atualize este guia quando houver novos serviços, canais, slogans ou padrões de campanha.")

doc.core_properties.title = "Guia de Marca — Mattos Soluções Digitais"
doc.core_properties.subject = "Identidade visual, aplicações e padrões de uso"
doc.core_properties.keywords = "Mattos, soluções digitais, marca, logo, Instagram, cartão de visita"
doc.core_properties.author = "Mattos Soluções Digitais"
doc.save(OUT)
print(OUT)
